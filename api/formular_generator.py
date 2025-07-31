"""
Formular‑Generator für VergabeAssist
-----------------------------------

Für die automatisierte Erstellung von Eigenerklärungen und Vertragsformularen
wird häufig auf PDF‑ oder DOCX‑Vorlagen zurückgegriffen.  Dieses Modul
stellt einfache Hilfsfunktionen bereit, um Platzhalter in Vorlagen
(z. B. ``{PROJEKTNAME}``, ``{AUFTRAGGEBER}``) durch tatsächliche
Projekt‑ und Auftraggeberdaten zu ersetzen.  Aktuell wird nur das
DOCX‑Format unterstützt, da PDF‑Formulare sehr unterschiedlich
strukturiert sein können und ein generisches Ausfüllen außerhalb des
Umfangs dieses Beispiels liegt.

Für PDF‑Formulare könnten Sie Bibliotheken wie ``pypdf`` oder
``pdfrw`` einsetzen und die Formularfelder programmatisch ausfüllen.
Die in diesem Modul enthaltenen Funktionen können als Ausgangspunkt
dienen, um eine solche PDF‑Implementierung zu ergänzen.
"""

from typing import Dict
from pathlib import Path
from docx import Document  # python‑docx


def fill_docx_template(template_path: str, output_path: str, context: Dict[str, str]) -> None:
    """Fülle Platzhalter in einer DOCX‑Vorlage aus und speichere das Ergebnis.

    :param template_path: Pfad zur DOCX‑Vorlage.  Platzhalter müssen in
        geschweiften Klammern angegeben werden, z. B. ``{PROJEKTNAME}``.
    :param output_path: Zielpfad für das ausgefüllte Dokument.
    :param context: Dictionary mit Schlüssel‑Wert‑Paaren, die die
        Platzhalterwerte definieren.  Beispielsweise ``{"PROJEKTNAME":
        "Sanierung Turnhalle"}``.
    """
    tpl_path = Path(template_path)
    if not tpl_path.exists():
        raise FileNotFoundError(f"Vorlage nicht gefunden: {template_path}")
    doc = Document(template_path)
    # Ersetze Platzhalter in allen Absätzen
    for paragraph in doc.paragraphs:
        for key, value in context.items():
            placeholder = f"{{{key}}}"
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)
    # Ersetze Platzhalter in Tabellenzellen
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in context.items():
                    placeholder = f"{{{key}}}"
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)
    # Speichere das ausgefüllte Dokument
    doc.save(output_path)


__all__ = ["fill_docx_template"]